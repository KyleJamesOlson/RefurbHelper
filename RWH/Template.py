# template.py
import re
import logging
from datetime import datetime
from docx import Document
from docx.shared import Pt
from utils import process_element

def fill_template(template_path, output_path, data, camera_found, form_data, keyname_fallback):
    doc = Document(template_path)

    current_date = datetime.now().strftime("%m/%d/%Y")
    filename_date = current_date.replace("/", "")

    logging.debug(f"Available sections in data: {list(data.keys())}")

    brand_section = next((key for key in data if 'Computer Brand Name' in data[key]), None)
    serial_section = next((key for key in data if 'Product Serial Number' in data[key]), None) or next((key for key in data if key == "System"), None)

    logging.debug(f"Brand section: {brand_section}")
    if brand_section:
        logging.debug(f"Brand section data: {data[brand_section]}")
    logging.debug(f"Serial section: {serial_section}")
    if serial_section:
        logging.debug(f"Serial section data: {data[serial_section]}")

    video_chipsets = keyname_fallback.get("Video Chipset", "N/A")
    # Fix: Clean video chipset strings by removing "Video Chipset:" prefix and normalizing spaces
    cleaned_chipsets = re.sub(r'^Video Chipset:\s*', '', video_chipsets) #:\s*', '', chipset).strip() for chipset in video_chipsets]
    video_chipset_str = "".join(cleaned_chipsets) if cleaned_chipsets else "N/A"
    logging.debug(f"Video chipset string: {video_chipset_str}")

    screen_size = ""
    resolution = ""
    monitor_section = next((key for key in data if 'Supported Video Modes' in data[key]), None)
    if not monitor_section:
        monitor_section = next((key for key in data if 'Monitor Name (Manuf)' in data[key]), None)
    logging.debug(f"Monitor section: {monitor_section}")
    if monitor_section:
        monitor_data = data[monitor_section]
        logging.debug(f"Monitor section data: {monitor_data}")
        if 'Monitor Name (Manuf)' in monitor_data:
            monitor_name = monitor_data['Monitor Name (Manuf)'].lower()
            if '17' in monitor_name:
                screen_size = "17\""
            elif '156' in monitor_name:
                screen_size = "15.6\""
            elif '154' in monitor_name:
                screen_size = "15.4\""
            elif '14' in monitor_name:
                screen_size = "14\""
            elif '13' in monitor_name:
                screen_size = "13\""
            elif '12' in monitor_name:
                screen_size = "12\""
            else:
                screen_size = "Unknown"
            logging.debug(f"Screen size determined: {screen_size}")
        if 'Supported Video Modes' in monitor_data:
            resolution = monitor_data['Supported Video Modes'].strip()
            logging.debug(f"Resolution: {resolution}")

    camera_present = 'Y' if camera_found else 'N'
    logging.debug(f"Camera present: {camera_present}")

    battery_info = "No"
    battery_section = next((key for key in data if 'Wear Level' in data[key]), None)
    if battery_section and 'Wear Level' in data[battery_section]:
        wear_level = float(data[battery_section]['Wear Level'].replace('%', ''))
        remaining_health = 100 - wear_level
        battery_info = f"Yes, {remaining_health:.1f}% remaining health"
        logging.debug(f"Battery info: {battery_info}")

    memory_info = ""
    memory_size = data.get("Memory", {}).get('Total Memory Size', '')
    memory_speed = data.get("Memory", {}).get('Memory Speed', '')

    logging.debug(f"Memory size: {memory_size}, Memory speed: {memory_speed}")
    if memory_size and memory_speed:
        ddr_match = re.search(r'DDR\d+-\d+', memory_speed)
        if ddr_match:
            speed_part = ddr_match.group(0)
            memory_info = f"{memory_size} {speed_part}MHz"
        else:
            memory_info = f"{memory_size}MHz"
        logging.debug(f"Memory info after parsing: {memory_info}")
    elif memory_size:
        memory_info = f"{memory_size}MHz"
    logging.debug(f"Final memory info: {memory_info}")

    network_info = "N/A"
    network_card = next((data[section].get('Network Card', '') for section in data if 'Network Card' in data[section]), '')
    link_speed = keyname_fallback.get("Maximum Link Speed", [])
    if network_card:
        network_info = f"{network_card} - {link_speed[0] if link_speed else '866 Mbps'}"
    logging.debug(f"Final network info: {network_info}")

    drive_model = "N/A"
    drive_section = next((key for key in data if 'Drive Model' in data[key]), None)
    if drive_section:
        drive_model = data[drive_section].get('Drive Model', 'N/A')
        logging.debug(f"Drive model: {drive_model}")

    dvd_cd = "None"

    audio_section = next((key for key in data if 'Audio Adapter' in data[key]), None)
    audio_adapter = data.get(audio_section, {}).get('Audio Adapter', 'N/A') if audio_section else 'N/A'
    logging.debug(f"Audio section: {audio_section}, Audio adapter: {audio_adapter}")

    os_value = keyname_fallback.get("Operating System", "N/A")
    os_cleaned = re.sub(r'^Operating System:\s*', '', os_value)
    logging.debug(f"Final operating system value: {os_cleaned}")

    power_adaptor = 'Yes' if form_data['power_adaptor'] else 'No'
    logging.debug(f"Power adaptor: {power_adaptor}")

    screen_label = "Touchscreen" if form_data['touchscreen'] else "Screen"

    brand_name = data.get(brand_section, {}).get('Computer Brand Name', 'N/A') if brand_section else 'N/A'
    serial_number = data.get(serial_section, {}).get('Product Serial Number', 'N/A') if serial_section else 'N/A'

    replacements = {
        '[1]': current_date,
        '[2]': form_data['technician_initials'],
        '[3]': brand_name,
        '[4]': serial_number,
        '[5]': data.get(serial_section, {}).get('SKU Number', 'N/A') if serial_section else 'N/A',
        '[6]': f"{'Yes' if form_data['warranty'] else 'No'}, {form_data['warranty_date'] if form_data['warranty'] else 'N/A'}",
        '[7]': data.get('Processor', {}).get('CPU Brand Name', 'N/A'),
        '[8]': memory_info if memory_info else 'N/A',
        '[9]': drive_model if drive_model != "N/A" else 'N/A',
        '[10]': dvd_cd,
        '[11]': network_info if network_info != "N/A" else 'N/A',
        '[12]': f"{screen_size} {screen_label}, {resolution} res, {video_chipset_str}",
        '[13]': audio_adapter,
        '[14]': os_cleaned if os_cleaned != "N/A" else "N/A",
        '[15]': battery_info,
        '[16]': power_adaptor,
        '[17]': camera_present,
        '[18]': form_data['ports'],
        '[19]': form_data['condition']
    }

    logging.debug(f"Replacements: {replacements}")

    for paragraph in doc.paragraphs:
        process_element(paragraph, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_element(paragraph, replacements)

    doc.save(output_path)
    logging.info(f"Document saved to {output_path}")